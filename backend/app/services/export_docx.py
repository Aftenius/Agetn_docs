import io
import re
from html import unescape

from bs4 import BeautifulSoup, Tag
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt


def _clean_text(text: str) -> str:
    text = unescape(text)
    return re.sub(r"\s+", " ", text).strip()


def _iter_block_nodes(root: Tag):
    for child in root.children:
        if not isinstance(child, Tag):
            continue
        if child.name == "div":
            yield from _iter_block_nodes(child)
        elif child.name in ("h1", "h2", "h3", "p", "ul", "ol", "table"):
            yield child


def html_to_docx_bytes(html: str) -> bytes:
    soup = BeautifulSoup(html, "html.parser")
    root: Tag = soup.body if soup.body else soup
    doc = Document()
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(14)

    nodes = list(_iter_block_nodes(root))
    if not nodes:
        text = _clean_text(soup.get_text())
        if text:
            doc.add_paragraph(text)
    else:
        for node in nodes:
            if node.name in ("h1", "h2", "h3"):
                level = int(node.name[1])
                p = doc.add_heading(_clean_text(node.get_text()), level=min(level, 3))
                p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                for run in p.runs:
                    run.font.name = "Times New Roman"
                try:
                    p.runs[0].font.size = Pt(14 if level >= 2 else 16)
                except IndexError:
                    pass
            elif node.name == "p":
                t = _clean_text(node.get_text())
                if t:
                    para = doc.add_paragraph(t)
                    for run in para.runs:
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(14)
            elif node.name in ("ul", "ol"):
                for li in node.find_all("li", recursive=False):
                    txt = _clean_text(li.get_text())
                    if txt:
                        para = doc.add_paragraph(txt, style="List Bullet")
                        for run in para.runs:
                            run.font.name = "Times New Roman"
                            run.font.size = Pt(14)
            elif node.name == "table":
                for row in node.find_all("tr"):
                    cells = [_clean_text(td.get_text()) for td in row.find_all(["td", "th"])]
                    line = " | ".join(c for c in cells if c)
                    if line:
                        para = doc.add_paragraph(line)
                        for run in para.runs:
                            run.font.name = "Times New Roman"
                            run.font.size = Pt(14)

    if len(doc.paragraphs) == 0:
        doc.add_paragraph("(Пустой документ)")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
